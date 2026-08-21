#!/usr/bin/env python3
"""Extract all text from an Excel workbook and consolidate into a Word document."""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


def cell_to_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def add_sheet_table(document: Document, sheet_name: str, rows: list[list[str]]) -> None:
    document.add_heading(sheet_name, level=1)

    if not rows:
        document.add_paragraph("(No data in this sheet.)")
        return

    max_cols = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]

    table = document.add_table(rows=len(normalized_rows), cols=max_cols)
    table.style = "Table Grid"

    for row_idx, row in enumerate(normalized_rows):
        for col_idx, text in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = text

    document.add_paragraph("")


def extract_workbook_to_docx(xlsx_path: Path, docx_path: Path) -> dict[str, int]:
    workbook = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True)
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    document.add_heading("HKTVMALL Ingredient Ranking — Extracted Text", level=0)
    document.add_paragraph(f"Source file: {xlsx_path.name}")
    document.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}"
    )
    document.add_paragraph(f"Sheets: {len(workbook.sheetnames)}")
    document.add_paragraph("")

    stats: dict[str, int] = {"sheets": 0, "rows": 0, "cells_with_text": 0}

    for sheet_name in workbook.sheetnames:
        worksheet = workbook[sheet_name]
        rows: list[list[str]] = []

        for row in worksheet.iter_rows(values_only=True):
            text_row = [cell_to_text(value) for value in row]
            while text_row and text_row[-1] == "":
                text_row.pop()
            if not any(text_row):
                continue
            rows.append(text_row)
            stats["rows"] += 1
            stats["cells_with_text"] += sum(1 for cell in text_row if cell)

        add_sheet_table(document, sheet_name, rows)
        stats["sheets"] += 1

        if sheet_name != workbook.sheetnames[-1]:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    workbook.close()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)
    return stats


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--input",
        type=Path,
        default=Path("HKTVMALL_Ingredient_Ranking.xlsx"),
        help="Path to the source Excel file",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("exports/HKTVMALL_Ingredient_Ranking_extracted.docx"),
        help="Path for the consolidated Word document",
    )
    args = parser.parse_args()

    if not args.input.exists():
        raise SystemExit(f"Input file not found: {args.input}")

    stats = extract_workbook_to_docx(args.input, args.output)
    print(f"Wrote: {args.output}")
    print(
        "Summary: "
        f"{stats['sheets']} sheets, "
        f"{stats['rows']} non-empty rows, "
        f"{stats['cells_with_text']} cells with text"
    )


if __name__ == "__main__":
    main()
