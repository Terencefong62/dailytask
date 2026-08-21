#!/usr/bin/env python3
"""Extract text from Illustrator (.ai) poster files via OCR and consolidate into Word."""

from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path

import pymupdf
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

DEFAULT_FILES = [
    "V&L  分页海报繁体-A2(1).ai",
    "V&L 总海报繁体A2-0317(1).ai",
    "VL 五张ol A2-0226(1).ai",
    "价值观5张olA2-0304(1).ai",
    "价值观海报5张繁体-A2-0317.ai",
]


def clean_ocr_text(text: str) -> str:
    text = text.replace("\x0c", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def ocr_page(page: pymupdf.Page, zoom: float, lang: str, psm: str) -> str:
    matrix = pymupdf.Matrix(zoom, zoom)
    pixmap = page.get_pixmap(matrix=matrix, alpha=False)

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as handle:
        image_path = Path(handle.name)

    try:
        pixmap.save(str(image_path))
        result = subprocess.run(
            [
                "tesseract",
                str(image_path),
                "stdout",
                "-l",
                lang,
                "--psm",
                psm,
            ],
            capture_output=True,
            text=True,
            check=True,
        )
    finally:
        image_path.unlink(missing_ok=True)

    if result.stderr:
        # Tesseract writes progress to stderr; ignore unless OCR failed.
        pass

    return clean_ocr_text(result.stdout)


def extract_file_text(
    file_path: Path,
    zoom: float,
    lang: str,
    psm: str,
) -> list[tuple[int, str]]:
    document = pymupdf.open(file_path)
    pages: list[tuple[int, str]] = []

    try:
        for index, page in enumerate(document, start=1):
            text = ocr_page(page, zoom=zoom, lang=lang, psm=psm)
            pages.append((index, text))
    finally:
        document.close()

    return pages


def build_word_document(
    output_path: Path,
    extracted: list[tuple[str, list[tuple[int, str]]]],
    source_dir: Path,
) -> dict[str, int]:
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    document.add_heading("V&L Poster Text Extraction", level=0)
    document.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
    document.add_paragraph(f"Source folder: {source_dir}")
    document.add_paragraph(f"Files processed: {len(extracted)}")
    document.add_paragraph("")

    stats = {"files": 0, "pages": 0, "pages_with_text": 0, "characters": 0}

    for file_index, (file_name, pages) in enumerate(extracted):
        document.add_heading(file_name, level=1)

        for page_number, text in pages:
            document.add_heading(f"Page {page_number}", level=2)
            if text:
                for paragraph in text.split("\n\n"):
                    paragraph = paragraph.strip()
                    if paragraph:
                        document.add_paragraph(paragraph)
                stats["pages_with_text"] += 1
                stats["characters"] += len(text)
            else:
                document.add_paragraph("(No text detected on this page.)")

            stats["pages"] += 1

        stats["files"] += 1
        if file_index < len(extracted) - 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return stats


def resolve_input_files(source_dir: Path, names: list[str]) -> list[Path]:
    files: list[Path] = []
    for name in names:
        path = source_dir / name
        if not path.exists():
            raise FileNotFoundError(f"Missing input file: {path}")
        files.append(path)
    return files


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source-dir",
        type=Path,
        default=Path("."),
        help="Directory containing the uploaded .ai files",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("exports/VL_posters_extracted_text.docx"),
        help="Output Word document path",
    )
    parser.add_argument(
        "--zoom",
        type=float,
        default=2.5,
        help="Render zoom factor for OCR (higher = sharper, slower)",
    )
    parser.add_argument(
        "--lang",
        default="chi_tra+eng",
        help="Tesseract language pack(s)",
    )
    parser.add_argument(
        "--psm",
        default="3",
        help="Tesseract page segmentation mode",
    )
    args = parser.parse_args()

    input_files = resolve_input_files(args.source_dir, DEFAULT_FILES)
    extracted: list[tuple[str, list[tuple[int, str]]]] = []

    for file_path in input_files:
        print(f"Processing {file_path.name} ...")
        pages = extract_file_text(
            file_path,
            zoom=args.zoom,
            lang=args.lang,
            psm=args.psm,
        )
        extracted.append((file_path.name, pages))
        non_empty = sum(1 for _, text in pages if text)
        print(f"  pages: {len(pages)}, with text: {non_empty}")

    stats = build_word_document(args.output, extracted, args.source_dir)
    print(f"Wrote: {args.output}")
    print(
        "Summary: "
        f"{stats['files']} files, "
        f"{stats['pages']} pages, "
        f"{stats['pages_with_text']} pages with text, "
        f"{stats['characters']} characters"
    )


if __name__ == "__main__":
    main()
